import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.oxml import OxmlElement
from io import BytesIO

# Function to add a horizontal line using a 1-cell table with bottom border
def add_horizontal_line(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.allow_autofit = True
    cell = table.cell(0, 0)
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.get_or_add_tcBorders()
    bottom = borders._element.xpath('./w:bottom')
    if not bottom:
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set('w:val', 'single')
        bottom_border.set('w:sz', '12')  # Thickness
        bottom_border.set('w:space', '0')
        bottom_border.set('w:color', '000000')
        borders._element.append(bottom_border)

# Function to generate Word bill
def generate_docx(customer_name, bill_to, bill_date, items):
    doc = Document()

    # Set 0.5 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Company Header
    doc.add_heading('S. S. ENTERPRISES', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(
        'Aluminium Interior Works\nPlot No. 651/A, East Kakatiyanagar, Neredmet, Malkajgiri, Secunderabad – 500056\n'
        'Cell: 9014462295, 7999110733'
    ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add bold full-width horizontal line
    add_horizontal_line(doc)

    # Customer name (left aligned)
    name_para = doc.add_paragraph(f"Customer Name: {customer_name}")
    name_para.runs[0].bold = True

    # Date (right aligned)
    date_para = doc.add_paragraph(f"Date: {bill_date}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_para.runs[0].bold = True

    # Bill to section
    doc.add_paragraph(f'Bill to: {bill_to}')
    doc.add_paragraph('BILL', style='Heading 2').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Table Headers
    headers = [
        "S.No.", "Item Name", "Sub Item Name", "Width in Sq.ft", "Height in Sq.ft", "Depth Sq.ft",
        "Total Sq.ft", "Price Sq.ft", "Per Sq.ft/Each", "Total Price"
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    # Table Rows
    for item in items:
        row = table.add_row().cells
        for i, h in enumerate(headers):
            row[i].text = str(item.get(h, ""))

    # Add Payment Terms section
    doc.add_paragraph("\nNote - Payment Terms and Conditions", style='Heading 2')

    terms = [
        "1. Advance Payment: An initial advance of 25% of the total project cost is required before commencement of work.",
        "2. Work Initiation: A further 25% is to be paid once the work has officially started.",
        "3. Post-Framing Stage: An additional 25% is to be paid upon completion of the framing stage.",
        "4. Final Payment: The remaining 25% must be paid upon completion of the finishing work."
    ]

    for term in terms:
        doc.add_paragraph(term, style='List Number')

    # Save Word to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Streamlit UI
st.set_page_config(page_title="S. S. Enterprises Bill Generator", layout="wide")
st.title("🧾 S. S. Enterprises - Bill Generator")

# Form inputs
customer_name = st.text_input("Customer Name", value="")
bill_to = st.text_input("Bill To", value="")
bill_date = st.date_input("Bill Date")

st.markdown("### Add Items Below")

item_count = st.number_input("How many items?", min_value=1, max_value=50, value=5)
items = []

for i in range(int(item_count)):
    with st.expander(f"Item #{i+1}"):
        col1, col2, col3 = st.columns(3)
        sno = col1.text_input("S.No.", value=str(i+1), key=f"sno_{i}")
        item_name = col2.text_input("Item Name", key=f"item_{i}")
        sub_item = col3.text_input("Sub Item Name", key=f"sub_{i}")

        col4, col5, col6 = st.columns(3)
        width = col4.text_input("Width in Sq.ft", key=f"width_{i}")
        height = col5.text_input("Height in Sq.ft", key=f"height_{i}")
        depth = col6.text_input("Depth Sq.ft", key=f"depth_{i}")

        col7, col8, col9 = st.columns(3)
        total_sqft = col7.text_input("Total Sq.ft", key=f"tsqft_{i}")
        price_per_sqft = col8.text_input("Price Sq.ft", key=f"pps_{i}")
        per_each = col9.text_input("Per Sq.ft/Each", key=f"per_{i}")

        total_price = st.text_input("Total Price", key=f"tprice_{i}")

        items.append({
            "S.No.": sno,
            "Item Name": item_name,
            "Sub Item Name": sub_item,
            "Width in Sq.ft": width,
            "Height in Sq.ft": height,
            "Depth Sq.ft": depth,
            "Total Sq.ft": total_sqft,
            "Price Sq.ft": price_per_sqft,
            "Per Sq.ft/Each": per_each,
            "Total Price": total_price
        })

# Generate Word file
if st.button("Generate Word Bill"):
    file = generate_docx(customer_name, bill_to, bill_date, items)

    # Clean filename using customer name
    safe_name = customer_name.strip().replace(" ", "_")
    filename = f"S_S_Enterprises_{safe_name}.docx"

    st.success("✅ Word document generated successfully!")
    st.download_button(
        "📥 Download Bill (.docx)",
        data=file,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
